import os
import io
import re
import time
import hashlib
import unicodedata
from typing import Union, Optional, Dict, Any
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor, as_completed

import pdfplumber
import docx
import filetype
from tqdm import tqdm  # For progress tracking

# Cache for parsed resumes to avoid re-processing
PARSED_RESUME_CACHE: Dict[str, str] = {}
CACHE_EXPIRY = 3600  # 1 hour cache expiry

# Thread pool for parallel processing
THREAD_POOL = ThreadPoolExecutor(max_workers=4)


def _clean_text(text: str) -> str:
    """Normalize and clean extracted text."""
    if not text:
        return ""
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


def _extract_pdf(file: Union[str, bytes]) -> str:
    """Extract text from PDF with optimizations."""
    cache_key = None
    if isinstance(file, str):
        cache_key = f"file:{os.path.abspath(file)}"
        # Check cache first
        if cache_key in PARSED_RESUME_CACHE:
            return PARSED_RESUME_CACHE[cache_key]
    
    try:
        start_time = time.time()
        
        # Optimized PDF extraction with page-level parallelization
        def process_page(page):
            try:
                return page.extract_text() or ""
            except Exception:
                return ""
        
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                with ThreadPoolExecutor() as executor:
                    pages = list(executor.map(process_page, pdf.pages))
        else:
            with pdfplumber.open(io.BytesIO(file)) as pdf:
                with ThreadPoolExecutor() as executor:
                    pages = list(executor.map(process_page, pdf.pages))
        
        text = "\n".join(pages)
        if not text.strip():
            return "No selectable text — OCR not supported in demo"
            
        result = _clean_text(text)
        
        # Cache the result
        if cache_key:
            PARSED_RESUME_CACHE[cache_key] = result
            
        return result
        
    except Exception as e:
        return f"Error parsing PDF: {str(e)[:200]}"  # Truncate long error messages


def _extract_docx(file: Union[str, bytes]) -> str:
    try:
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(io.BytesIO(file))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return _clean_text("\n".join(paragraphs))
    except Exception as e:
        return f"Error parsing DOCX: {e}"


def _extract_txt(file: Union[str, bytes]) -> str:
    try:
        if isinstance(file, str):
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            text = file.decode("utf-8", errors="ignore")
        return _clean_text(text)
    except Exception as e:
        return f"Error parsing TXT: {e}"


def extract_text_from_file(file: Union[str, bytes], use_cache: bool = True) -> str:
    """
    Extract plain text from a resume file (PDF, DOCX, TXT, MD).
    Returns cleaned plain text.
    """
    if isinstance(file, str):  # path
        ext = os.path.splitext(file)[-1].lower()
        if ext == ".pdf":
            return _extract_pdf(file)
        elif ext == ".docx":
            return _extract_docx(file)
        elif ext in [".txt", ".md"]:
            # Generate cache key for bytes input
            cache_key = None
            if use_cache and isinstance(file, str) and os.path.isfile(file):
                file_stat = os.stat(file)
                cache_key = f"{file}:{file_stat.st_mtime}"
                if cache_key in PARSED_RESUME_CACHE:
                    return PARSED_RESUME_CACHE[cache_key]
                with open(file, 'rb') as f:
                    file_bytes = f.read()
            else:
                file_bytes = file if isinstance(file, bytes) else file.read()
                if use_cache:
                    cache_key = f"bytes:{hashlib.md5(file_bytes).hexdigest()}"
                    if cache_key in PARSED_RESUME_CACHE:
                        return PARSED_RESUME_CACHE[cache_key]

            # Detect file type
            kind = filetype.guess(file_bytes)

            mime = kind.mime if kind else None
            if mime and "pdf" in mime:
                return _extract_pdf(file)
            elif mime and ("word" in mime or "officedocument" in mime):
                return _extract_docx(file)
            elif mime and "text" in mime:
                return _extract_txt(file)
            else:
                return "Unsupported file type"
        else:
            with open(file, "rb") as f:
                header = f.read(261)
            kind = filetype.guess(header)
    else:  # raw bytes
        kind = filetype.guess(file[:261])

    mime = kind.mime if kind else None
    if mime and "pdf" in mime:
        return _extract_pdf(file)
    elif mime and ("word" in mime or "officedocument" in mime):
        return _extract_docx(file)
    elif mime and "text" in mime:
        return _extract_txt(file)
    else:
        return "Unsupported file type"


def parse_folder(folder_path: str, output_file: str = "parsed_resumes.txt"):
    """Parse all resumes in a folder and save into one text file."""
    results = []
    for fname in os.listdir(folder_path):
        fpath = os.path.join(folder_path, fname)
        if os.path.isfile(fpath):
            text = extract_text_from_file(fpath)
            results.append(f"\n===== {fname} =====\n{text}\n")

    with open(output_file, "w", encoding="utf-8") as f:
        f.writelines(results)

    print(f"✅ Parsed {len(results)} resumes. Output saved to {output_file}")


if __name__ == "__main__":
    # Default: parse resumes folder (relative to backend/ingestion)
    folder = os.path.join("..", "..", "data", "resumes")
    if os.path.isdir(folder):
        parse_folder(folder, "parsed_resumes.txt")
    else:
        print("⚠️ Folder ../data/resumes not found.")
def parse_resume_file(file_path: str) -> str:
    """Return plain text from PDF/DOCX/TXT resume."""
    import subprocess
    import os

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return _extract_docx(file_path)
    elif ext in [".txt", ".md"]:
        return _extract_txt(file_path)
    else:
        return "Unsupported file type"
